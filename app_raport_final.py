import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import xlsxwriter

# ==========================================
# 1. KONFIGURASI & HELPER FUNCTIONS
# ==========================================
st.set_page_config(page_title="Sistem Raport SMA Islam Al-Ghozali", layout="wide", page_icon="🏫")

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def terbilang(n):
    angka = ["", "Satu", "Dua", "Tiga", "Empat", "Lima", "Enam", "Tujuh", "Delapan", "Sembilan", "Sepuluh", "Sebelas"]
    if n < 0 or n > 100: return ""
    elif n < 12: return angka[n]
    elif n < 20: return angka[n-10] + " Belas"
    elif n < 100: return angka[n//10] + " Puluh " + angka[n%10]
    elif n == 100: return "Seratus"
    return ""

# ==========================================
# 2. INISIALISASI DATABASE
# ==========================================
if 'sekolah_info' not in st.session_state:
    st.session_state['sekolah_info'] = {
        "nama": "SMA ISLAM AL-GHOZALI",
        "alamat": "Jl. Permata No. 19 Desa Curug Kec. Gunungsindur Kab. Bogor Telp. (0251)8614072",
        "kepsek": "Antoni Firdaus M.Pd.",
        "semester": "Genap", 
        "tahun_ajar": "2024/2025",
        "kota_raport": "Gunungsindur",
        "tgl_raport": "20 Maret 2025"
    }

# Data Master Global
if 'master_kelas' not in st.session_state: st.session_state['master_kelas'] = ["X-A", "X-B", "XI-A", "XII-A"]
if 'master_guru' not in st.session_state: st.session_state['master_guru'] = ["Liyas Syarifudin, M.Pd", "Mali, S.Pd", "Antoni Firdaus M.Pd."]
if 'master_mapel' not in st.session_state:
    st.session_state['master_mapel'] = ["Pendidikan Agama Islam", "Pendidikan Pancasila", "Bahasa Indonesia", "Matematika", "Fisika (IPA)", "Kimia (IPA)", "Biologi (IPA)", "Sosiologi", "Ekonomi", "Sejarah", "Geografi", "Bahasa Inggris", "PJOK", "Informatika", "Seni Budaya", "Prakarya", "Bahasa Sunda"]

# DATA TRANSAKSIONAL
# Format data_kkm: Key string "KELAS|MAPEL" -> Value int KKM
# Contoh: "X-A|Matematika": 75
if 'data_kkm' not in st.session_state: st.session_state['data_kkm'] = {}

# Struktur Mapel per Kelas (Untuk Admin menentukan mapel apa saja di kelas tsb)
# Format: {'X-A': ['Matematika', 'Fisika'], ...}
if 'mapel_per_kelas' not in st.session_state: st.session_state['mapel_per_kelas'] = {}

if 'penugasan_guru' not in st.session_state: st.session_state['penugasan_guru'] = [] 
if 'data_wali_kelas' not in st.session_state: st.session_state['data_wali_kelas'] = {} 
if 'data_siswa' not in st.session_state: st.session_state['data_siswa'] = [] 
if 'data_nilai' not in st.session_state: st.session_state['data_nilai'] = [] 
if 'data_non_akademik' not in st.session_state: st.session_state['data_non_akademik'] = {} 

if 'login_status' not in st.session_state: st.session_state['login_status'] = False; st.session_state['user_role'] = None

# ==========================================
# 3. LOGIC
# ==========================================
def set_kkm_guru(kelas, mapel, nilai_kkm):
    """Guru menyimpan KKM untuk kelas dan mapel tertentu"""
    key = f"{kelas}|{mapel}"
    st.session_state['data_kkm'][key] = nilai_kkm

def get_kkm_final(kelas, mapel):
    """Mengambil KKM. Prioritas: Input Guru > Default 75"""
    key = f"{kelas}|{mapel}"
    return st.session_state['data_kkm'].get(key, 75)

def get_list_mapel_raport(kelas):
    """Mengambil daftar mapel yang aktif di kelas tersebut (Settingan Admin)"""
    # Jika admin sudah set spesifik, pakai itu. Jika belum, pakai semua master mapel.
    return st.session_state['mapel_per_kelas'].get(kelas, st.session_state['master_mapel'])

def get_nilai(sid, mapel):
    for d in st.session_state['data_nilai']:
        if d['siswa_id'] == sid and d['mapel'] == mapel: return d['nilai']
    return 0

def update_nilai(sid, mapel, val):
    found = False
    for d in st.session_state['data_nilai']:
        if d['siswa_id'] == sid and d['mapel'] == mapel:
            d['nilai'] = val; found = True; break
    if not found:
        st.session_state['data_nilai'].append({"siswa_id": sid, "mapel": mapel, "nilai": val})

def klaim_pengajaran(guru, mapel, kelas):
    found = False
    for p in st.session_state['penugasan_guru']:
        if p['mapel'] == mapel and p['kelas'] == kelas:
            p['nama'] = guru; found = True; break
    if not found:
        st.session_state['penugasan_guru'].append({"nama": guru, "mapel": mapel, "kelas": kelas})

def get_guru_pengajar(mapel, kelas):
    for p in st.session_state['penugasan_guru']:
        if p['mapel'] == mapel and p['kelas'] == kelas: return p['nama']
    return None

def update_non_akademik(sid, data): st.session_state['data_non_akademik'][sid] = data
def get_non_akademik(sid): return st.session_state['data_non_akademik'].get(sid, {"kerapihan": "-", "kedisiplinan": "-", "kejujuran": "-", "sakit": 0, "izin": 0, "alpha": 0})

def hitung_ranking(kelas):
    siswa_kelas = [s for s in st.session_state['data_siswa'] if s['kelas'] == kelas]
    rekap = []
    for s in siswa_kelas:
        total = 0
        for d in st.session_state['data_nilai']:
            if d['siswa_id'] == s['id']: total += d['nilai']
        rekap.append({"id": s['id'], "total": total})
    rekap_sorted = sorted(rekap, key=lambda x: x['total'], reverse=True)
    rank_dict = {item['id']: idx + 1 for idx, item in enumerate(rekap_sorted)}
    return rank_dict, len(siswa_kelas)

def logout(): st.session_state['login_status'] = False; st.rerun()

# ==========================================
# 4. WORD GENERATOR
# ==========================================
def generate_docx_custom(siswa, rank, total_siswa):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5); section.right_margin = Inches(0.5)

    sek = st.session_state['sekolah_info']
    nama_wali = st.session_state['data_wali_kelas'].get(siswa['kelas'], "(............)")
    data_pribadi = get_non_akademik(siswa['id'])

    def add_centered(txt, sz=12, b=True):
        p = doc.add_paragraph(txt); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold=b; p.runs[0].font.size=Pt(sz); return p

    add_centered("LAPORAN HASIL BELAJAR PESERTA DIDIK")
    add_centered("PENILAIAN TENGAH SEMESTER GENAP", 14)
    add_centered(sek['nama'], 16)
    add_centered(f"TAHUN PELAJARAN {sek['tahun_ajar']}")
    p_al = add_centered(sek['alamat'], 10, False)
    
    p_bdr = p_al._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr'); btm = OxmlElement('w:bottom')
    btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '12'); bdr.append(btm); p_bdr.append(bdr)
    doc.add_paragraph()

    # Identitas
    nipd_disp = siswa.get('nipd','-'); nisn_disp = siswa.get('nisn','-')
    ti = doc.add_table(3,4); ti.autofit=False; ti.columns[0].width=Inches(1.5)
    r0=ti.rows[0].cells; r0[0].text="Nama"; r0[1].text=f": {siswa['nama']}"; r0[2].text="NIPD"; r0[3].text=f": {nipd_disp}"
    r1=ti.rows[1].cells; r1[0].text="Sekolah"; r1[1].text=f": {sek['nama']}"; r1[2].text="NISN"; r1[3].text=f": {nisn_disp}"
    r2=ti.rows[2].cells; r2[0].text="Kls/Smt"; r2[1].text=f": {siswa['kelas']}/{sek['semester']}"; r2[2].text="Thn"; r2[3].text=f": {sek['tahun_ajar']}"
    doc.add_paragraph()

    # Nilai
    tn = doc.add_table(2,6); tn.style='Table Grid'
    h0=tn.rows[0].cells; h1=tn.rows[1].cells
    h0[0].merge(h1[0]).text="NO"; h0[1].merge(h1[1]).text="Komponen Mata Pelajaran"
    h0[2].merge(h1[2]).text="KKM"; h0[3].merge(h0[4]).text="Nilai"
    h1[3].text="Angka"; h1[4].text="Huruf"; h0[5].merge(h1[5]).text="Afektif"

    for c in h0+h1: 
        set_cell_bg(c, "E0F7FA")
        if c.paragraphs:
            p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True

    tot=0; cnt=0
    # Ambil mapel sesuai kelas
    mapel_raport = get_list_mapel_raport(siswa['kelas'])
    
    for idx, m in enumerate(mapel_raport):
        r = tn.add_row().cells
        val = get_nilai(siswa['id'], m)
        kkm = get_kkm_final(siswa['kelas'], m) # AMBIL KKM DARI INPUT GURU
        
        r[0].text=str(idx+1); r[1].text=m; r[2].text=str(kkm); r[3].text=str(val); r[4].text=terbilang(val)
        
        # Logika Predikat (Sesuai KKM)
        afektif = "B" if val >= kkm else "C" if val > 0 else "-"
        r[5].text=afektif
        
        for cell in r: 
            if cell.paragraphs: cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        if r[1].paragraphs: r[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        tot+=val; 
        if val>0: cnt+=1

    rs = tn.add_row().cells; rs[0].merge(rs[2]).text="Jumlah"; rs[3].text=str(tot)
    if rs[0].paragraphs: rs[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; rs[0].paragraphs[0].runs[0].bold=True
    if rs[3].paragraphs: rs[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    ra = tn.add_row().cells; ra[0].merge(ra[2]).text="Rata - rata"; ra[3].text=f"{tot/cnt:.2f}" if cnt>0 else "0"
    if ra[0].paragraphs: ra[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; ra[0].paragraphs[0].runs[0].bold=True
    if ra[3].paragraphs: ra[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Kepribadian & Absen
    doc.add_paragraph("B. Kepribadian dan Ketidakhadiran").runs[0].bold=True
    tc = doc.add_table(1,2); tc.style='Table Grid'
    tk = tc.cell(0,0).add_table(4,2); tk.style='Table Grid'
    tk.cell(0,0).text="Kepribadian"; tk.cell(0,1).text="Ket"
    tk.cell(1,0).text="Kerapihan"; tk.cell(1,1).text=str(data_pribadi['kerapihan'])
    tk.cell(2,0).text="Kedisiplinan"; tk.cell(2,1).text=str(data_pribadi['kedisiplinan'])
    tk.cell(3,0).text="Kejujuran"; tk.cell(3,1).text=str(data_pribadi['kejujuran'])
    
    ta = tc.cell(0,1).add_table(4,2); ta.style='Table Grid'
    ta.cell(0,0).text="Absensi"; ta.cell(0,1).text="Ket"
    ta.cell(1,0).text="Sakit"; ta.cell(1,1).text=str(data_pribadi['sakit'])
    ta.cell(2,0).text="Izin"; ta.cell(2,1).text=str(data_pribadi['izin'])
    ta.cell(3,0).text="Alpha"; ta.cell(3,1).text=str(data_pribadi['alpha'])

    doc.add_paragraph(f"\nPeringkat Kelas: {rank} dari {total_siswa} siswa")
    ttd = doc.add_table(1,3); ttd.alignment=WD_TABLE_ALIGNMENT.CENTER
    c1=ttd.cell(0,0); c2=ttd.cell(0,1); c3=ttd.cell(0,2)
    c1.text="\nOrang Tua\n\n\n(..........)"
    c2.text=f"\nMengetahui\nKepala Sekolah\n\n\n({sek['kepsek']})"
    c3.text=f"{sek['kota_raport']}, {sek['tgl_raport']}\nWali Kelas\n\n\n({nama_wali})"
    for c in ttd.rows[0].cells: 
        if c.paragraphs: c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio

# ==========================================
# 5. HALAMAN ADMIN
# ==========================================
def admin_page():
    st.sidebar.title("Admin Panel")
    menu = st.sidebar.radio("Navigasi", ["🏠 Dashboard", "📚 Atur Mapel per Kelas", "👨‍🎓 Data Siswa", "⚙️ Data Master", "📊 Monitoring", "👨‍🏫 Atur Wali Kelas", "⚙️ Info Sekolah"])
    st.title("Administrator")

    # --- ATUR MAPEL PER KELAS ---
    if menu == "📚 Atur Mapel per Kelas":
        st.subheader("Struktur Kurikulum")
        st.info("Pilih mapel apa saja yang muncul di Raport untuk kelas tertentu.")
        
        col1, col2 = st.columns(2)
        with col1:
            pilih_kelas = st.selectbox("Pilih Kelas", st.session_state['master_kelas'])
        
        # Ambil mapel yg sudah ada
        current_mapel = st.session_state['mapel_per_kelas'].get(pilih_kelas, st.session_state['master_mapel'])
        
        with col2:
            new_mapels = st.multiselect("Mapel Aktif", st.session_state['master_mapel'], default=current_mapel)
            
        if st.button("Simpan Struktur Kurikulum"):
            st.session_state['mapel_per_kelas'][pilih_kelas] = new_mapels
            st.success(f"Berhasil update kurikulum {pilih_kelas}!")

    # --- DATA SISWA (Copy Paste) ---
    elif menu == "👨‍🎓 Data Siswa":
        t1, t2 = st.tabs(["📋 Copy-Paste Excel", "Lihat Data"])
        with t1:
            st.info("Format: **KELAS | Nama | NIPD | JK | NISN**")
            raw = st.text_area("Paste Data Siswa", height=200)
            if st.button("Proses Simpan"):
                lines = raw.strip().split('\n'); cnt=0
                for line in lines:
                    parts = line.split('\t') if '\t' in line else line.split(',')
                    parts = [p.strip() for p in parts]
                    if len(parts) >= 2:
                        k = parts[0]; n = parts[1]
                        nipd = parts[2] if len(parts)>2 else "-"
                        jk = parts[3] if len(parts)>3 else "-"
                        nisn = parts[4] if len(parts)>4 else "-"
                        if k.upper() != "KELAS":
                            if k not in st.session_state['master_kelas']: st.session_state['master_kelas'].append(k)
                            st.session_state['data_siswa'].append({"id":f"C{len(st.session_state['data_siswa'])}","nama":n,"kelas":k,"nisn":nisn,"nipd":nipd,"jk":jk})
                            cnt+=1
                st.success(f"Masuk {cnt} siswa")
        with t2: st.dataframe(pd.DataFrame(st.session_state['data_siswa']))

    # --- DATA MASTER ---
    elif menu == "⚙️ Data Master":
        t1, t2, t3 = st.tabs(["Guru", "Mapel", "Kelas"])
        with t1:
            raw = st.text_area("Paste Guru", height=150)
            if st.button("Simpan Guru"):
                for l in raw.split('\n'):
                    if l.strip() and l.strip() not in st.session_state['master_guru']: st.session_state['master_guru'].append(l.strip())
                st.success("Updated")
        with t2:
            raw = st.text_area("Paste Mapel", height=150)
            if st.button("Simpan Mapel"):
                for l in raw.split('\n'):
                    if l.strip() and l.strip() not in st.session_state['master_mapel']: st.session_state['master_mapel'].append(l.strip())
                st.success("Updated")
        with t3:
            raw = st.text_area("Paste Kelas", height=150)
            if st.button("Simpan Kelas"):
                for l in raw.split('\n'):
                    if l.strip() and l.strip() not in st.session_state['master_kelas']: st.session_state['master_kelas'].append(l.strip())
                st.success("Updated")

    elif menu == "📊 Monitoring":
        st.subheader("Monitoring Input Nilai")
        # Menampilkan juga KKM yang sudah diinput guru
        data = []
        for m in st.session_state['master_mapel']:
            row = {"Mapel": m}
            for k in st.session_state['master_kelas']:
                pg = get_guru_pengajar(m, k)
                kkm = get_kkm_final(k, m)
                status = f"✅ {pg} (KKM:{kkm})" if pg else "❌"
                row[k] = status
            data.append(row)
        st.dataframe(pd.DataFrame(data), use_container_width=True)

    elif menu == "👨‍🏫 Atur Wali Kelas":
        with st.form("wali"):
            k=st.selectbox("Kelas",st.session_state['master_kelas']); g=st.selectbox("Wali",st.session_state['master_guru'])
            if st.form_submit_button("Set"): st.session_state['data_wali_kelas'][k]=g; st.success("Ok")
        st.write(st.session_state['data_wali_kelas'])

    elif menu == "⚙️ Info Sekolah":
        with st.form("sch"):
            i=st.session_state['sekolah_info']
            n=st.text_input("Nama",i['nama']); a=st.text_input("Alamat",i['alamat']); k=st.text_input("Kepsek",i['kepsek'])
            ci=st.text_input("Kota",i['kota_raport']); tg=st.text_input("Tgl",i['tgl_raport'])
            if st.form_submit_button("Simpan"): st.session_state['sekolah_info'].update({"nama":n,"alamat":a,"kepsek":k,"kota_raport":ci,"tgl_raport":tg}); st.success("Ok")

    elif menu == "🏠 Dashboard":
        st.metric("Total Siswa", len(st.session_state['data_siswa']))

# ==========================================
# 6. HALAMAN GURU
# ==========================================
def guru_page():
    g=st.session_state['active_user']; k=st.session_state['active_kelas']; m=st.session_state['active_mapel']
    st.sidebar.title("Panel Guru")
    st.sidebar.info(f"👨‍🏫 {g}\n📚 {m}\n🏫 {k}")
    st.title(f"Input Nilai")
    
    sis=[s for s in st.session_state['data_siswa'] if s['kelas']==k]
    if not sis: st.warning("Kelas Kosong"); return
    
    # === INPUT KKM OLEH GURU ===
    st.markdown("### 1. Tentukan KKM")
    current_kkm = get_kkm_final(k, m)
    input_kkm = st.number_input("KKM Mata Pelajaran Ini:", min_value=0, max_value=100, value=current_kkm)
    
    st.markdown("### 2. Input Nilai Siswa")
    
    t1, t2, t3 = st.tabs(["📝 Manual", "📂 Upload File", "📋 Copy-Paste"])
    
    with t1:
        with st.form("input_manual"):
            tmp={}
            c1,c2,c3=st.columns([1,4,2]); c1.write("No"); c2.write("Nama"); c3.write("Nilai")
            for i,s in enumerate(sis):
                ca,cb,cc=st.columns([1,4,2]); ca.write(f"{i+1}"); cb.write(s['nama'])
                tmp[s['id']] = cc.number_input(f"v_{s['id']}", 0, 100, int(get_nilai(s['id'], m)), label_visibility='collapsed')
            
            if st.form_submit_button("💾 Simpan Data"):
                for sid,v in tmp.items(): update_nilai(sid,m,v)
                # Simpan KKM
                set_kkm_guru(k, m, input_kkm)
                # Klaim Pengajaran
                klaim_pengajaran(g, m, k)
                st.success("Tersimpan!")

    with t2:
        st.info("Format: **Nama | Nilai**")
        f = st.file_uploader("Upload Excel/CSV", type=['xlsx','csv'])
        if f and st.button("Proses Upload"):
            try:
                df = pd.read_csv(f) if f.name.endswith('.csv') else pd.read_excel(f)
                cnt=0
                for i,r in df.iterrows():
                    r={k.lower():v for k,v in r.items()}
                    nm=r.get('nama',''); val=r.get('nilai',0)
                    for s in sis:
                        if nm.lower() in s['nama'].lower():
                            update_nilai(s['id'], m, val); cnt+=1; break
                set_kkm_guru(k, m, input_kkm)
                klaim_pengajaran(g, m, k)
                st.success(f"Updated {cnt} siswa")
            except Exception as e: st.error(f"Error: {e}")

    with t3:
        st.info("Copy kolom **Nama** dan **Nilai** dari Excel")
        raw = st.text_area("Paste Data", height=200)
        if st.button("Proses Paste"):
            cnt=0
            for l in raw.split('\n'):
                p = l.split('\t') if '\t' in l else l.split(',')
                if len(p)>=2:
                    nm=p[0].strip(); 
                    try: val=int(float(p[1].strip()))
                    except: val=0
                    for s in sis:
                        if nm.lower() in s['nama'].lower():
                            update_nilai(s['id'],m,val); cnt+=1; break
            set_kkm_guru(k, m, input_kkm)
            klaim_pengajaran(g, m, k)
            st.success(f"Updated {cnt} siswa")

# ==========================================
# 7. HALAMAN WALI KELAS
# ==========================================
def wali_page():
    w=st.session_state['active_user']; k=st.session_state['active_kelas']
    st.title(f"Wali Kelas: {k}"); sis=[s for s in st.session_state['data_siswa'] if s['kelas']==k]
    rank_dict, tot_sis = hitung_ranking(k)
    
    t1,t2,t3 = st.tabs(["Non-Akademik", "Leger", "Raport"])
    with t1:
        with st.form("non"):
            tmp={}
            for s in sis:
                with st.expander(s['nama']):
                    c1,c2=st.columns(2); old=get_non_akademik(s['id'])
                    with c1:
                        kr=st.selectbox(f"Rapi {s['id']}",["-","AA","BB","CC"],index=["-","AA","BB","CC"].index(old.get('kerapihan','-')) if old.get('kerapihan') in ["-","AA","BB","CC"] else 0)
                        kd=st.selectbox(f"Disiplin {s['id']}",["-","AA","BB","CC"],index=["-","AA","BB","CC"].index(old.get('kedisiplinan','-')) if old.get('kedisiplinan') in ["-","AA","BB","CC"] else 0)
                        kj=st.selectbox(f"Jujur {s['id']}",["-","AA","BB","CC"],index=["-","AA","BB","CC"].index(old.get('kejujuran','-')) if old.get('kejujuran') in ["-","AA","BB","CC"] else 0)
                    with c2:
                        sa=st.number_input(f"Sakit {s['id']}",0,100,old.get('sakit',0))
                        iz=st.number_input(f"Izin {s['id']}",0,100,old.get('izin',0))
                        al=st.number_input(f"Alpha {s['id']}",0,100,old.get('alpha',0))
                    tmp[s['id']]={"kerapihan":kr,"kedisiplinan":kd,"kejujuran":kj,"sakit":sa,"izin":iz,"alpha":al}
            if st.form_submit_button("Simpan"):
                for sid,d in tmp.items(): update_non_akademik(sid,d)
                st.success("Ok")
    with t2:
        # Tampilkan mapel sesuai kurikulum kelas
        mapel_kelas = get_list_mapel_raport(k)
        shorts={m:m[:4].upper() for m in mapel_kelas}
        data=[]
        for i,s in enumerate(sis):
            r={"No":i+1,"Nama":s['nama'], "NISN":s.get('nisn','-')}
            tot=0; cnt=0
            for m in mapel_kelas:
                val=get_nilai(s['id'],m); r[shorts[m]]=val; tot+=val; 
                if val>0: cnt+=1
            non=get_non_akademik(s['id'])
            r.update({"Rata":f"{tot/cnt:.2f}" if cnt else "0", "Rank":rank_dict.get(s['id']), "S":non.get('sakit'), "I":non.get('izin'), "A":non.get('alpha')})
            data.append(r)
        
        df=pd.DataFrame(data)
        st.dataframe(df, use_container_width=True, hide_index=True)
        buf=io.BytesIO()
        with pd.ExcelWriter(buf, engine='xlsxwriter') as w: df.to_excel(w, sheet_name='Leger', index=False)
        st.download_button("Download Excel", buf, f"Leger_{k}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    with t3:
        for s in sis:
            c1,c2=st.columns([4,1]); c1.write(f"**{s['nama']}** (Rank {rank_dict.get(s['id'])})")
            docx = generate_docx_custom(s, rank_dict.get(s['id']), tot_sis)
            c2.download_button("Download", docx, f"Raport_{s['nama']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{s['id']}")

# ==========================================
# 8. LOGIN SCREEN
# ==========================================
def login_screen():
    st.title("PORTAL RAPORT AL-GHOZALI"); t1,t2,t3=st.tabs(["ADMIN","WALI","GURU"])
    
    with t1:
        if st.button("Log Admin") and st.text_input("Pw",type="password")=="admin": st.session_state['login_status']=True; st.session_state['user_role']='admin'; st.rerun()
    
    with t2:
        w=st.selectbox("Nama Wali",st.session_state['master_guru'],key='wn'); k=st.selectbox("Kelas Binaan",st.session_state['master_kelas'],key='wk')
        if st.button("Log Wali"): st.session_state['login_status']=True; st.session_state['user_role']='wali'; st.session_state['active_user']=w; st.session_state['active_kelas']=k; st.rerun()
    
    with t3:
        gn = st.selectbox("Nama Guru", st.session_state['master_guru'], key='gn')
        gm = st.selectbox("Mata Pelajaran", st.session_state['master_mapel'], key='gm')
        gk = st.selectbox("Kelas Ajar", st.session_state['master_kelas'], key='gk')
        
        if st.button("Mulai Input Nilai"):
            st.session_state['login_status']=True; st.session_state['user_role']='guru'; st.session_state['active_user']=gn; st.session_state['active_kelas']=gk; st.session_state['active_mapel']=gm; st.rerun()

if not st.session_state['login_status']: login_screen()
else:
    with st.sidebar:
        if st.button("Logout"): logout()
    r=st.session_state['user_role']
    if r=='admin': admin_page()
    elif r=='guru': guru_page()
    elif r=='wali': wali_page()